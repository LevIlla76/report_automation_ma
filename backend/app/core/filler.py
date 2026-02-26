# backend/core/filler.py

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class DocumentFiller:
    def __init__(self, template_path):
        self.doc = Document(template_path)

    def find_column_index(self, table_idx, header_row_idx, keywords):
        """ ค้นหา Index แบบเดิม (ต้องเจอทุกคำในช่องเดียว) """
        try:
            table = self.doc.tables[table_idx]
            header_cells = table.rows[header_row_idx].cells
            for i, cell in enumerate(header_cells):
                cell_text = cell.text.lower().replace("\n", " ").replace("\r", " ")
                if all(k.lower() in cell_text for k in keywords):
                    return i
            return -1
        except:
            return -1

    def find_nth_column_index(self, table_idx, header_row_idx, keyword, n=1):
        """ 
        ค้นหา keyword ว่าอยู่ที่คอลัมน์ไหน โดยระบุลำดับได้ (n)
        เช่น n=1 คือเจอครั้งแรก (AVG), n=2 คือเจอครั้งที่สอง (MAX)
        """
        try:
            table = self.doc.tables[table_idx]
            header_cells = table.rows[header_row_idx].cells
            count = 0
            for i, cell in enumerate(header_cells):
                cell_text = cell.text.lower().replace("\n", " ")
                if keyword.lower() in cell_text:
                    count += 1
                    if count == n:
                        return i
            return -1
        except:
            return -1
        
    def replace_text_with_image(self, text_placeholder, image_path, width=Pt(400)):
        """ 
        ค้นหาข้อความ placeholder (เช่น {{image_f5_internet}}) 
        แล้วลบข้อความนั้นทิ้ง ก่อนจะแทรกรูปภาพลงไปแทนที่ 
        """
        # 1. ค้นหาในย่อหน้าปกติ (นอกตาราง)
        for p in self.doc.paragraphs:
            if text_placeholder in p.text:
                p.text = p.text.replace(text_placeholder, "")
                run = p.add_run()
                run.add_picture(image_path, width=width)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. ค้นหาในตาราง
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if text_placeholder in p.text:
                            p.text = p.text.replace(text_placeholder, "")
                            run = p.add_run()
                            run.add_picture(image_path, width=width)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER    

    def fill_slot(self, table_idx, row_idx, col_idx, value, is_image=False):
        try:
            table = self.doc.tables[table_idx]
            
            # --- Safety Check: ห้ามเขียนทับ Header (แถว 0) เด็ดขาด ---
            if row_idx == 0:
                print(f"⚠️ Attempted to write to Header (Row 0). Skipped. Value: {value}")
                return

            if row_idx >= len(table.rows) or col_idx >= len(table.rows[row_idx].cells):
                return

            cell = table.rows[row_idx].cells[col_idx]
            cell.text = "" 
            paragraph = cell.paragraphs[0]
            
            if is_image and value and os.path.exists(str(value)):
                run = paragraph.add_run()
                run.add_picture(value, width=Pt(150))
            else:
                run = paragraph.add_run(str(value))
                run.font.name = 'TH Sarabun New'
                run.font.size = Pt(14)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        except Exception as e:
            print(f"Fill error at T{table_idx} R{row_idx} C{col_idx}: {e}")

    def save(self, output_path):
        self.doc.save(output_path)